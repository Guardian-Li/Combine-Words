from docx import Document
import os
from docxcompose.composer import Composer
from docx import Document as Document_compose

path="word"
folder = os.path.exists(path)
if not folder:
    os.makedirs(path)
    print("mk new folder:word")


origin_word=[]

file_list = os.listdir("word")
print(file_list)
origin_word_list=os.listdir("word/"+file_list[0])
stu={}
for x in origin_word_list:
    file_name = "word/"+file_list[0]+"/"+x
    #merge_document=Document()
    master=Document(file_name)
    master.add_page_break()
    composer=Composer(master)
    stu[x]=composer
    #stu[x].add_page_break()



for x in range(1,len(file_list)):
    for y in os.listdir("word/"+file_list[x]):
        if y in stu:
            # sub_doc = Document()
            # if x < len(file_list) - 1:
            #      sub_doc.add_page_break()
            # sub_doc.save(y)
            #
            # for element in sub_doc.element.body:
            #     stu[y].element.body.append(element)
            file_name = "word/"+file_list[x]+"/"+y

            doc2= Document(file_name)
            #doc2.add_page_break()
            stu[y].append(doc2)

for key in stu:
    print(key)
    stu[key].save(key)





